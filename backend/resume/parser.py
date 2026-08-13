"""
AI Job Hunter - Document Parser

Extracts raw text from PDF and DOCX resume files.
"""

from __future__ import annotations

from pathlib import Path

from backend.utils.logger import get_logger

logger = get_logger("resume.parser")

ALLOWED_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
}


def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Extracted text content.
    """
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())

    full_text = "\n\n".join(pages)
    logger.info("PDF text extracted", pages=len(reader.pages), chars=len(full_text))
    return full_text


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text content.
    """
    from docx import Document

    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text += "\n" + row_text

    logger.info("DOCX text extracted", paragraphs=len(paragraphs), chars=len(full_text))
    return full_text


def extract_text(file_path: str, content_type: str) -> str:
    """Extract text from a resume file based on its content type.

    Args:
        file_path: Path to the file.
        content_type: MIME type of the file.

    Returns:
        Extracted text content.

    Raises:
        ValueError: If the file type is not supported.
    """
    if content_type == "application/pdf":
        return extract_text_from_pdf(file_path)
    elif content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/docx",
    ):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {content_type}")


def validate_file_type(content_type: str, filename: str) -> str:
    """Validate that the uploaded file is a supported resume format.

    Args:
        content_type: MIME type from the upload.
        filename: Original filename.

    Returns:
        The validated content type.

    Raises:
        ValueError: If the file type is not supported.
    """
    # Check by MIME type first
    if content_type in ALLOWED_TYPES:
        return content_type

    # Fallback: check by extension
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "application/pdf"
    elif ext == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    raise ValueError(
        f"Unsupported file type '{content_type}' for file '{filename}'. "
        "Only PDF and DOCX files are supported."
    )
